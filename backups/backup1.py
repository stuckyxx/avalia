import streamlit as st
import json
import os
from datetime import datetime, timedelta
import docx
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_ORIENT
from docx2pdf import convert

# --- LISTA DE MUNICÍPIOS ---
MUNICIPIOS_MARANHAO = [
    "Açailândia", "Afonso Cunha", "Água Doce do Maranhão", "Alcântara", "Aldeias Altas", 
    "Altamira do Maranhão", "Alto Alegre do Maranhão", "Alto Alegre do Pindaré", "Alto Parnaíba", 
    "Amapá do Maranhão", "Amarante do Maranhão", "Anajatuba", "Anapurus", "Apicum-Açu", 
    "Araguanã", "Araioses", "Arame", "Arari", "Axixá", "Bacabal", "Bacabeira", "Bacuri", 
    "Bacurituba", "Balsas", "Barão de Grajaú", "Barra do Corda", "Barreirinhas", 
    "Bela Vista do Maranhão", "Belágua", "Benedito Leite", "Bequimão", "Bernardo do Mearim", 
    "Boa Vista do Gurupi", "Bom Jardim", "Bom Jesus das Selvas", "Bom Lugar", "Brejo", 
    "Brejo de Areia", "Buriti", "Buriti Bravo", "Buriticupu", "Buritirana", "Cachoeira Grande", 
    "Cajapió", "Cajari", "Campestre do Maranhão", "Cândido Mendes", "Cantanhede", 
    "Capinzal do Norte", "Carolina", "Carutapera", "Caxias", "Cedral", "Central do Maranhão", 
    "Centro do Guilherme", "Centro Novo do Maranhão", "Chapadinha", "Cidelândia", "Codó", 
    "Coelho Neto", "Colinas", "Conceição do Lago-Açu", "Coroatá", "Cururupu", "Davinópolis", 
    "Dom Pedro", "Duque Bacelar", "Esperantinópolis", "Estreito", "Feira Nova do Maranhão", 
    "Fernando Falcão", "Formosa da Serra Negra", "Fortaleza dos Nogueiras", "Fortuna", 
    "Godofredo Viana", "Gonçalves Dias", "Governador Archer", "Governador Edison Lobão", 
    "Governador Eugênio Barros", "Governador Luiz Rocha", "Governador Newton Bello", 
    "Governador Nunes Freire", "Graça Aranha", "Grajaú", "Guimarães", "Humberto de Campos", 
    "Icatu", "Igarapé do Meio", "Igarapé Grande", "Imperatriz", "Itaipava do Grajaú", 
    "Itapecuru Mirim", "Itinga do Maranhão", "Jatobá", "Jenipapo dos Vieiras", "João Lisboa", 
    "Joselândia", "Junco do Maranhão", "Lago da Pedra", "Lago do Junco", "Lago dos Rodrigues", 
    "Lago Verde", "Lagoa do Mato", "Lagoa Grande do Maranhão", "Lajeado Novo", "Lima Campos", 
    "Loreto", "Luís Domingues", "Magalhães de Almeida", "Maracaçumé", "Marajá do Sena", 
    "Maranhãozinho", "Mata Roma", "Matinha", "Matões", "Matões do Norte", "Milagres do Maranhão", 
    "Mirador", "Miranda do Norte", "Mirinzal", "Monção", "Montes Altos", "Morros", "Nina Rodrigues", 
    "Nova Colinas", "Nova Iorque", "Nova Olinda do Maranhão", "Olho d'Água das Cunhãs", 
    "Olinda Nova do Maranhão", "Paço do Lumiar", "Palmeirândia", "Paraibano", "Parnarama", 
    "Passagem Franca", "Pastos Bons", "Paulino Neves", "Paulo Ramos", "Pedreiras", 
    "Pedro do Rosário", "Penalva", "Peri Mirim", "Peritoró", "Pindaré-Mirim", "Pinheiro", 
    "Pio XII", "Pirapemas", "Poção de Pedras", "Porto Franco", "Porto Rico do Maranhão", 
    "Presidente Dutra", "Presidente Juscelino", "Presidente Médici", "Presidente Sarney", 
    "Presidente Vargas", "Primeira Cruz", "Raposa", "Riachão", "Ribamar Fiquene", "Rosário", 
    "Sambaíba", "Santa Filomena do Maranhão", "Santa Helena", "Santa Inês", "Santa Luzia", 
    "Santa Luzia do Paruá", "Santa Quitéria do Maranhão", "Santa Rita", "Santana do Maranhão", 
    "Santo Amaro do Maranhão", "Santo Antônio dos Lopes", "São Benedito do Rio Preto", 
    "São Bento", "São Bernardo", "São Domingos do Azeitão", "São Domingos do Maranhão", 
    "São Félix de Balsas", "São Francisco do Brejão", "São Francisco do Maranhão", 
    "São João Batista", "São João do Carú", "São João do Paraíso", "São João do Soter", 
    "São João dos Patos", "São José de Ribamar", "São José dos Basílios", "São Luís", 
    "São Luís Gonzaga do Maranhão", "São Mateus do Maranhão", "São Pedro da Água Branca", 
    "São Pedro dos Crentes", "São Raimundo das Mangabeiras", "São Raimundo do Doca Bezerra", 
    "São Roberto", "São Vicente Ferrer", "Satubinha", "Senador Alexandre Costa", 
    "Senador La Rocque", "Serrano do Maranhão", "Sítio Novo", "Sucupira do Norte", 
    "Sucupira do Riachão", "Tasso Fragoso", "Timbiras", "Timon", "Trizidela do Vale", 
    "Tufilândia", "Tuntum", "Turiaçu", "Turilândia", "Tutóia", "Urbano Santos", 
    "Vargem Grande", "Viana", "Vila Nova dos Martírios", "Vitória do Mearim", "Vitorino Freire", "Zé Doca"
]

# --- FUNÇÕES AUXILIARES ---
@st.cache_data
def carregar_criterios_do_arquivo(caminho_arquivo="criterios_por_topico.json"):
    try:
        with open(caminho_arquivo, 'r', encoding='utf-8') as f: return json.load(f)
    except FileNotFoundError: st.error(f"ERRO: O arquivo de dados '{caminho_arquivo}' não foi encontrado."); return None
    except json.JSONDecodeError: st.error(f"ERRO: O arquivo '{caminho_arquivo}' contém um erro de formatação."); return None

def criar_pastas_necessarias():
    os.makedirs("data/avaliacoes", exist_ok=True)
    os.makedirs("relatorios", exist_ok=True)

def calcular_indices_por_secao(respostas, matriz_perguntas):
    pesos = {"ESSENCIAL": 2.0, "OBRIGATÓRIA": 1.5, "RECOMENDADA": 1.0}
    indices_secao = {}
    for secao, perguntas in matriz_perguntas.items():
        pontos_possiveis_secao = 0
        pontos_obtidos_secao = 0
        for item in perguntas:
            classificacao = item.get("classificacao", "RECOMENDADA").upper()
            peso = pesos.get(classificacao, 1.0)
            pontos_possiveis_secao += peso
            status_geral_atende = not any(respostas.get(f"{secao}_{item['criterio']}_{sub}") == "Não Atende" for sub in item["subcriterios"])
            if status_geral_atende:
                pontos_obtidos_secao += peso
        indice = (pontos_obtidos_secao / pontos_possiveis_secao * 100) if pontos_possiveis_secao > 0 else 0
        indices_secao[secao] = indice
    return indices_secao

def calcular_indice_e_selo(respostas, matriz_perguntas):
    pesos = {"ESSENCIAL": 2.0, "OBRIGATÓRIA": 1.5, "RECOMENDADA": 1.0}
    total_pontos_possiveis, pontos_obtidos, total_essenciais, essenciais_atendidos = 0, 0, 0, 0
    for secao, perguntas in matriz_perguntas.items():
        for item in perguntas:
            classificacao = item.get("classificacao", "RECOMENDADA").upper()
            peso = pesos.get(classificacao, 1.0)
            total_pontos_possiveis += peso
            status_geral_atende = not any(respostas.get(f"{secao}_{item['criterio']}_{sub}") == "Não Atende" for sub in item["subcriterios"])
            if status_geral_atende: pontos_obtidos += peso
            if classificacao == "ESSENCIAL":
                total_essenciais += 1
                if status_geral_atende: essenciais_atendidos += 1
    percentual_essenciais = (essenciais_atendidos / total_essenciais * 100) if total_essenciais > 0 else 100
    indice = (pontos_obtidos / total_pontos_possiveis * 100) if total_pontos_possiveis > 0 else 0
    selo = "Inexistente"
    if indice > 0:
        if percentual_essenciais == 100:
            if indice >= 95: selo = "💎 Diamante"
            elif indice >= 85: selo = "🥇 Ouro"
            elif indice >= 75: selo = "🥈 Prata"
            else: selo = "Elevado (não elegível para selo)"
        else:
            if indice >= 75: selo = "Elevado"
            elif indice >= 50: selo = "Intermediário"
            elif indice >= 30: selo = "Básico"
            else: selo = "Inicial"
    return {"indice": indice, "selo": selo, "percentual_essenciais": percentual_essenciais}

def gerar_relatorio_detalhado_pdf(respostas, municipio, segmento, matriz_perguntas, tipo_relatorio):
    doc = docx.Document()
    
    doc.add_heading('PROGRAMA NACIONAL DE TRANSPARÊNCIA PÚBLICA', level=1).alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_heading(f'Relatório de Transparência\n{segmento} de {municipio}', level=2).alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph()
    resultados = calcular_indice_e_selo(respostas, matriz_perguntas)
    p_resultado = doc.add_paragraph(); p_resultado.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_indice = p_resultado.add_run(f"{resultados['indice']:.2f}%\n"); run_indice.font.size = Pt(20); run_indice.bold = True
    run_selo = p_resultado.add_run(f"{resultados['selo']}"); run_selo.font.size = Pt(16); run_selo.bold = True
    doc.add_paragraph()
    doc.add_paragraph(f"Exercício: {datetime.now().year}")
    doc.add_paragraph(f"Avaliação feita por: Avaliador do Sistema")
    doc.add_paragraph(f"Data de Geração: {datetime.now().strftime('%d/%m/%Y %H:%M:%S')}")
    
    doc.add_page_break()
    doc.add_heading('Detalhamento da Avaliação', level=1)

    indices_por_secao = calcular_indices_por_secao(respostas, matriz_perguntas)
    
    matriz_a_usar = matriz_perguntas
    if tipo_relatorio == "Apenas Não Conformidades":
        perguntas_filtradas = {}
        for secao, perguntas in matriz_perguntas.items():
            itens_nao_conformes = [item for item in perguntas if any(respostas.get(f"{secao}_{item['criterio']}_{sub}") == "Não Atende" for sub in item["subcriterios"])]
            if itens_nao_conformes:
                perguntas_filtradas[secao] = itens_nao_conformes
        matriz_a_usar = perguntas_filtradas

    if not matriz_a_usar:
        doc.add_paragraph("Nenhuma não conformidade foi encontrada.")
    else:
        for secao, perguntas in matriz_a_usar.items():
            percentual_secao = indices_por_secao.get(secao, 0)
            doc.add_heading(f"{secao} - {percentual_secao:.2f}%", level=2)
            for item in perguntas:
                doc.add_heading(f"{item.get('topico', '')} - {item.get('criterio', '')} ({item.get('classificacao', '')})", level=3)
                for subcriterio in item["subcriterios"]:
                    p_sub = doc.add_paragraph(style='List Bullet')
                    status = respostas.get(f"{secao}_{item['criterio']}_{subcriterio}", "Não Avaliado")
                    p_sub.add_run(f"{subcriterio}: ")
                    run_status = p_sub.add_run(status)
                    run_status.font.color.rgb = RGBColor(0x00, 0x80, 0x00) if status == "Atende" else RGBColor(0xFF, 0x00, 0x00)
                    run_status.bold = True
                p_evidencias = doc.add_paragraph(); p_evidencias.add_run("\nEvidências e Comentários:").bold = True
                chave_links_pergunta = f"{secao}_{item['criterio']}_links"
                links = respostas.get(chave_links_pergunta, [])
                for link in links:
                    doc.add_paragraph(f"- Link: {link}")
                for subcriterio in item["subcriterios"]:
                     obs = respostas.get(f"{secao}_{item['criterio']}_{subcriterio}_obs", "")
                     if obs:
                         doc.add_paragraph(f"- Observação ({subcriterio}): {obs}")
                doc.add_paragraph()
    
    timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
    prefixo = "NaoConformidades" if tipo_relatorio == "Apenas Não Conformidades" else "Completo"
    nome_base = f"Relatorio_Detalhado_{prefixo}_{segmento.replace(' ', '')}_{municipio.replace(' ', '')}_{timestamp}"
    path_docx = os.path.join("relatorios", f"{nome_base}.docx")
    path_pdf = os.path.join("relatorios", f"{nome_base}.pdf")
    doc.save(path_docx)
    try:
        convert(path_docx, path_pdf); os.remove(path_docx)
        return path_pdf
    except Exception as e:
        st.sidebar.error(f"Falha ao converter para PDF. Verifique se o MS Word está instalado.")
        st.sidebar.info("O arquivo Word (.docx) foi gerado como alternativa.")
        st.session_state.fallback_docx_path = path_pdf
        return None

def on_disponibilidade_change(secao, criterio, subcriterios):
    chave_disponibilidade = f"{secao}_{criterio}_Disponibilidade"
    novo_status_disponibilidade = st.session_state.respostas[chave_disponibilidade]
    novo_status_subs = "Atende" if novo_status_disponibilidade == "Atende" else "Não Atende"
    for sub in subcriterios:
        if sub != "Disponibilidade": st.session_state.respostas[f"{secao}_{criterio}_{sub}"] = novo_status_subs

# --- INTERFACE GRÁFICA ---
st.set_page_config(layout="wide", page_title="Avaliador de Transparência")
st.title("📄 Sistema de Avaliação de Transparência Municipal")
st.markdown("---")
matriz_completa = carregar_criterios_do_arquivo()

if matriz_completa:
    criar_pastas_necessarias()
    st.sidebar.header("Configuração da Avaliação")
    opcoes_municipios = ["- Selecione um município -"] + MUNICIPIOS_MARANHAO
    municipio = st.sidebar.selectbox("Nome do Município", options=opcoes_municipios)
    segmento = st.sidebar.selectbox("Órgão/Poder", list(matriz_completa.keys()))
    if municipio != "- Selecione um município -" and segmento:
        nome_arquivo_avaliacao = f"avaliacao_{segmento.replace(' ', '')}_{municipio.replace(' ', '')}.json"
        caminho_arquivo = os.path.join("data/avaliacoes", nome_arquivo_avaliacao)
        if st.sidebar.button("✅ Iniciar / Continuar Avaliação"):
            if os.path.exists(caminho_arquivo):
                with open(caminho_arquivo, 'r', encoding='utf-8') as f: st.session_state.respostas = json.load(f)
                st.sidebar.success("Avaliação anterior carregada!")
            else:
                st.session_state.respostas = {}; st.sidebar.info("Iniciando uma nova avaliação.")
            st.session_state.path_pdf = None; st.session_state.fallback_docx_path = None
            st.session_state.avaliacao_iniciada = True; st.session_state.caminho_arquivo = caminho_arquivo
            st.session_state.municipio = municipio; st.session_state.segmento = segmento
            st.session_state.last_save_time = datetime.now()
            
    if st.session_state.get('avaliacao_iniciada', False):
        if 'last_save_time' not in st.session_state: st.session_state.last_save_time = datetime.now()
        if datetime.now() - st.session_state.last_save_time > timedelta(minutes=10):
            try:
                with open(st.session_state.caminho_arquivo, 'w', encoding='utf-8') as f: json.dump(st.session_state.respostas, f, ensure_ascii=False, indent=4)
                st.session_state.last_save_time = datetime.now()
                st.toast(f"Progresso salvo automaticamente às {datetime.now().strftime('%H:%M:%S')}")
            except Exception as e:
                st.toast(f"Erro no salvamento automático: {e}")
        st.header(f"Avaliação: {st.session_state.municipio} - {st.session_state.segmento}")
        matriz_perguntas = matriz_completa[st.session_state.segmento]
        for secao, perguntas in matriz_perguntas.items():
            with st.expander(f"**{secao}**", expanded=False):
                for item in perguntas:
                    st.markdown(f"#### {item['topico']} - {item['criterio']}"); st.markdown("---")
                    col_link_ui, _ = st.columns([1, 1])
                    with col_link_ui:
                        st.subheader("Links de Evidência")
                        chave_links = f"{secao}_{item['criterio']}_links"
                        if chave_links not in st.session_state.respostas: st.session_state.respostas[chave_links] = []
                        for i, link in enumerate(st.session_state.respostas[chave_links]):
                            link_cols = st.columns([10, 1]); link_cols[0].info(link)
                            if link_cols[1].button("X", key=f"rem_{chave_links}_{i}"): st.session_state.respostas[chave_links].pop(i); st.rerun()
                        link_cols = st.columns([10, 1])
                        novo_link = link_cols[0].text_input("Adicionar novo link", key=f"add_{chave_links}", label_visibility="collapsed")
                        if link_cols[1].button("➕", key=f"btn_{chave_links}"):
                            if novo_link: st.session_state.respostas[chave_links].append(novo_link); st.rerun()
                    st.markdown("---"); st.subheader("Critérios de Avaliação")
                    subcriterios = item["subcriterios"]
                    if "Disponibilidade" in subcriterios:
                        cols = st.columns([1, 2]); subcriterio = "Disponibilidade"; chave_resposta = f"{secao}_{item['criterio']}_{subcriterio}"
                        with cols[0]:
                            resposta_atual = st.session_state.respostas.get(chave_resposta, "Atende")
                            resposta = st.radio(subcriterio, ("Atende", "Não Atende"), index=1 if resposta_atual == "Não Atende" else 0, key=chave_resposta, horizontal=True, on_change=on_disponibilidade_change, kwargs=dict(secao=secao, criterio=item['criterio'], subcriterios=subcriterios))
                            st.session_state.respostas[chave_resposta] = resposta
                        if resposta == "Não Atende":
                            with cols[1]:
                                chave_obs = f"{chave_resposta}_obs"; obs = st.text_area("Observação:", value=st.session_state.respostas.get(chave_obs, ""), key=chave_obs, height=10); st.session_state.respostas[chave_obs] = obs
                    chave_disponibilidade = f"{secao}_{item['criterio']}_Disponibilidade"
                    disponibilidade_falhou = ("Disponibilidade" in subcriterios and st.session_state.respostas.get(chave_disponibilidade) == "Não Atende")
                    if not disponibilidade_falhou:
                        for subcriterio in subcriterios:
                            if subcriterio != "Disponibilidade":
                                cols = st.columns([1, 2]); chave_resposta = f"{secao}_{item['criterio']}_{subcriterio}"
                                with cols[0]:
                                    resposta_atual = st.session_state.respostas.get(chave_resposta, "Atende"); resposta = st.radio(subcriterio, ("Atende", "Não Atende"), index=1 if resposta_atual == "Não Atende" else 0, key=chave_resposta, horizontal=True); st.session_state.respostas[chave_resposta] = resposta
                                if resposta == "Não Atende":
                                    with cols[1]:
                                        chave_obs = f"{chave_resposta}_obs"; obs = st.text_area("Observação:", value=st.session_state.respostas.get(chave_obs, ""), key=chave_obs, height=10); st.session_state.respostas[chave_obs] = obs
                    st.markdown("---")
        st.sidebar.header("Ações")
        if st.sidebar.button("💾 Salvar Progresso"):
            with open(st.session_state.caminho_arquivo, 'w', encoding='utf-8') as f: json.dump(st.session_state.respostas, f, ensure_ascii=False, indent=4)
            st.session_state.last_save_time = datetime.now(); st.sidebar.success("Progresso salvo!")
        st.sidebar.markdown("##### Tipo de Relatório"); tipo_relatorio = st.sidebar.radio("Escolha o tipo:", ("Apenas Não Conformidades", "Relatório Completo"), label_visibility="collapsed")
        if st.sidebar.button("📊 Gerar Relatório PDF"):
            with st.spinner("Gerando relatório PDF..."):
                with open(st.session_state.caminho_arquivo, 'w', encoding='utf-8') as f: json.dump(st.session_state.respostas, f, ensure_ascii=False, indent=4)
                st.session_state.fallback_docx_path = None
                path_pdf = gerar_relatorio_detalhado_pdf(st.session_state.respostas, st.session_state.municipio, st.session_state.segmento, matriz_completa[st.session_state.segmento], tipo_relatorio)
                st.session_state.path_pdf = path_pdf
            if st.session_state.path_pdf: st.sidebar.success("Relatório PDF pronto!")
        if st.session_state.get('path_pdf'):
            with open(st.session_state.path_pdf, "rb") as pdf_file: st.sidebar.download_button(label="⬇️ Baixar Relatório (.pdf)", data=pdf_file, file_name=os.path.basename(st.session_state.path_pdf), mime="application/pdf", key="download_pdf")
        if st.session_state.get('fallback_docx_path'):
            with open(st.session_state.fallback_docx_path, "rb") as docx_file: st.sidebar.download_button(label="⬇️ Baixar Arquivo Word (.docx)", data=docx_file, file_name=os.path.basename(st.session_state.fallback_docx_path), mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document", key="download_docx_fallback")
else:
    st.warning("Aguardando o carregamento do arquivo 'criterios_por_topico.json'...")