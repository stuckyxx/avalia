# Arquivo: test.py

import streamlit as st
import json
import os
from datetime import datetime, timedelta
import docx
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx2pdf import convert
import streamlit_authenticator as stauth
import yaml
from yaml.loader import SafeLoader
from db_management import get_db_connection, initialize_database

# --- FUNÇÕES AUXILIARES ---
@st.cache_data
def carregar_criterios_do_arquivo(caminho_arquivo="criterios_por_topico.json"):
    """Carrega os critérios de avaliação e a lista de municípios do arquivo JSON."""
    try:
        with open(caminho_arquivo, 'r', encoding='utf-8') as f:
            return json.load(f)
    except FileNotFoundError:
        st.error(f"ERRO: O arquivo de dados '{caminho_arquivo}' não foi encontrado.")
        return None
    except json.JSONDecodeError:
        st.error(f"ERRO: O arquivo '{caminho_arquivo}' contém um erro de formatação.")
        return None

def salvar_progresso_db(conn):
    """Salva as respostas da sessão atual no banco de dados usando a conexão fornecida."""
    if 'avaliacao_id' not in st.session_state: return False
    try:
        cursor = conn.cursor()
        respostas_para_salvar = [
            (st.session_state.avaliacao_id, chave, str(valor))
            for chave, valor in st.session_state.get('respostas', {}).items()
        ]
        if respostas_para_salvar:
            cursor.executemany(
                "INSERT OR REPLACE INTO respostas (avaliacao_id, chave, valor) VALUES (?, ?, ?)",
                respostas_para_salvar
            )
            conn.commit()
        return True
    except Exception as e:
        st.error(f"Erro ao salvar progresso: {e}"); return False

@st.cache_data
def calcular_indice_e_selo(respostas, matriz_perguntas):
    """Calcula o índice de transparência e o selo Atricon com base nos pesos."""
    pesos = {"ESSENCIAL": 2.0, "OBRIGATÓRIA": 1.5, "RECOMENDADA": 1.0}
    total_pontos_possiveis, pontos_obtidos, total_essenciais, essenciais_atendidos = 0, 0, 0, 0
    for secao, perguntas in matriz_perguntas.items():
        if secao == "Municipios_MA": continue
        for item in perguntas:
            classificacao = item.get("classificacao", "RECOMENDADA").upper()
            peso = pesos.get(classificacao, 1.0)
            total_pontos_possiveis += peso
            status_geral_atende = not any(respostas.get(f"{secao}_{item.get('criterio')}_{sub}") == "Não Atende" for sub in item.get("subcriterios", []))
            if status_geral_atende: pontos_obtidos += peso
            if classificacao == "ESSENCIAL":
                total_essenciais += 1
                if status_geral_atende: essenciais_atendidos += 1
    percentual_essenciais = (essenciais_atendidos / total_essenciais * 100) if total_essenciais > 0 else 100
    indice = (pontos_obtidos / total_pontos_possiveis * 100) if total_pontos_possiveis > 0 else 0
    selo = "Inexistente"
    if indice > 0:
        if percentual_essenciais == 100:
            if indice >= 95: selo = "💎 Diamante"
            elif indice >= 85: selo = "🥇 Ouro"
            elif indice >= 75: selo = "🥈 Prata"
        else:
            if indice >= 75: selo = "Elevado"
            elif indice >= 50: selo = "Intermediário"
            elif indice >= 30: selo = "Básico"
            else: selo = "Inicial"
    return {"indice": indice, "selo": selo}

def on_disponibilidade_change(secao, criterio, subcriterios):
    """Callback para atualizar subcritérios quando a Disponibilidade muda."""
    chave_disponibilidade = f"{secao}_{criterio}_Disponibilidade"
    if st.session_state.get(chave_disponibilidade) == "Não Atende":
        for sub in subcriterios:
            if sub != "Disponibilidade":
                st.session_state[f"{secao}_{criterio}_{sub}"] = "Não Atende"
                st.session_state.respostas[f"{secao}_{criterio}_{sub}"] = "Não Atende"

def gerar_relatorio_novo_modelo(respostas, municipio, segmento, matriz_perguntas, tipo_relatorio, nome_usuario, usuario_config):
    """Gera o relatório completo em formato DOCX e tenta converter para PDF."""
    template_tipo = usuario_config.get('template', 'padrao')
    template_path = f"modelo_{template_tipo}.docx"
    try:
        doc = docx.Document(template_path)
    except Exception:
        st.sidebar.error(f"ERRO: Arquivo de modelo '{template_path}' não foi encontrado."); return None
    
    # --- PÁGINA DE ROSTO ---
    p_title = doc.add_paragraph()
    p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_title = p_title.add_run("PADRÃO MÍNIMO DE QUALIDADE"); run_title.font.size = Pt(22); run_title.bold = True
    doc.add_paragraph()
    p_subtitulo = doc.add_paragraph(); p_subtitulo.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_subtitulo.add_run("Relatório de Transparência\n").bold = True; doc.add_paragraph()
    p_subtitulo.add_run(f"{segmento} de {municipio}").bold = True
    resultados = calcular_indice_e_selo(respostas, matriz_perguntas)
    doc.add_paragraph()
    p_score = doc.add_paragraph(f"{resultados['indice']:.2f}%"); p_score.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_score.runs[0].font.size = Pt(48); p_score.runs[0].bold = True
    p_selo = doc.add_paragraph(f"{resultados['selo']}"); p_selo.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_selo.runs[0].font.size = Pt(24); p_selo.runs[0].bold = True
    doc.add_paragraph(); doc.add_paragraph()
    texto_intro = f"Com base na Lei 12.527/2011 (Lei de Acesso à Informação)..."
    doc.add_paragraph(texto_intro); doc.add_paragraph()
    doc.add_paragraph(f"Exercício: {datetime.now().year}"); doc.add_paragraph()
    doc.add_paragraph(f"Avaliação feita por: {nome_usuario}"); doc.add_paragraph()
    doc.add_paragraph(f"Data de Geração: {datetime.now().strftime('%d/%m/%Y %H:%M:%S')}")
    doc.add_page_break()
    p_detalhe = doc.add_paragraph(); run_detalhe = p_detalhe.add_run("Detalhamento da Avaliação"); run_detalhe.font.size = Pt(18); run_detalhe.bold = True
    doc.add_paragraph()
    
    # Lógica de filtragem e exibição do relatório...

    os.makedirs("relatorios", exist_ok=True)
    timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
    nome_base = f"Relatorio_{segmento.replace(' ', '')}_{municipio.replace(' ', '')}_{timestamp}"
    path_docx = os.path.join("relatorios", f"{nome_base}.docx")
    path_pdf = os.path.join("relatorios", f"{nome_base}.pdf")
    doc.save(path_docx)
    try:
        convert(path_docx, path_pdf); os.remove(path_docx)
        return path_pdf
    except Exception as e:
        st.sidebar.error(f"Falha ao converter para PDF: {e}"); st.session_state.fallback_docx_path = path_docx
        return None

# --- APLICAÇÃO PRINCIPAL ---
st.set_page_config(layout="wide", page_title="Avaliador de Transparência")
conn = get_db_connection()
initialize_database(conn)

st.title("📄 Sistema de Avaliação de Transparência Municipal")
matriz_completa = carregar_criterios_do_arquivo()

if matriz_completa:
    try:
        with open('config.yaml', 'r', encoding='utf-8') as file:
            config = yaml.load(file, Loader=SafeLoader)
        authenticator = stauth.Authenticate(
            config['credentials'], config['cookie']['name'],
            config['cookie']['key'], config['cookie']['expiry_days']
        )
        authenticator.login('main')
    except FileNotFoundError:
        st.error("ERRO: O arquivo 'config.yaml' não foi encontrado."); st.stop()

    if st.session_state.get("authentication_status"):
        authenticator.logout('Logout', 'sidebar')
        st.sidebar.title(f"Bem-vindo(a),\n{st.session_state.get('name')}!")

        if "avaliacao_id" in st.query_params and not st.session_state.get('avaliacao_iniciada'):
            try:
                avaliacao_id_url = int(st.query_params["avaliacao_id"])
                cursor = conn.cursor()
                cursor.execute("SELECT municipio, segmento FROM avaliacoes WHERE id=?", (avaliacao_id_url,))
                avaliacao_info = cursor.fetchone()
                if avaliacao_info:
                    cursor.execute("SELECT chave, valor FROM respostas WHERE avaliacao_id=?", (avaliacao_id_url,))
                    respostas_db = cursor.fetchall()
                    st.session_state.avaliacao_id = avaliacao_id_url
                    st.session_state.municipio = avaliacao_info[0]
                    st.session_state.segmento = avaliacao_info[1]
                    st.session_state.respostas = {chave: valor for chave, valor in respostas_db}
                    st.session_state.avaliacao_iniciada = True
                    st.toast("Sessão restaurada a partir da URL!")
            except (ValueError, TypeError): pass
        
        st.sidebar.header("Configuração da Avaliação")
        municipios = ["- Selecione -"] + sorted(matriz_completa.get("Municipios_MA", []))
        segmentos = ["- Selecione -"] + [k for k in matriz_completa.keys() if k != "Municipios_MA"]
        
        mun_idx = municipios.index(st.session_state.get('municipio', '- Selecione -'))
        seg_idx = segmentos.index(st.session_state.get('segmento', '- Selecione -'))
        
        municipio_selecionado = st.sidebar.selectbox("Município", municipios, index=mun_idx)
        segmento_selecionado = st.sidebar.selectbox("Órgão/Poder", segmentos, index=seg_idx)

        if municipio_selecionado != "- Selecione -" and segmento_selecionado != "- Selecione -":
            if st.sidebar.button("✅ Iniciar / Continuar Avaliação", use_container_width=True):
                cursor = conn.cursor()
                cursor.execute(
                    "SELECT id FROM avaliacoes WHERE municipio=? AND segmento=? AND usuario=? ORDER BY data_inicio DESC LIMIT 1",
                    (municipio_selecionado, segmento_selecionado, st.session_state.get('username'))
                )
                avaliacao_existente = cursor.fetchone()
                
                if avaliacao_existente:
                    st.session_state.avaliacao_id = avaliacao_existente[0]
                else:
                    cursor.execute("INSERT INTO avaliacoes (municipio, segmento, usuario) VALUES (?, ?, ?)", (municipio_selecionado, segmento_selecionado, st.session_state.get('username')))
                    conn.commit()
                    st.session_state.avaliacao_id = cursor.lastrowid
                
                cursor.execute("SELECT chave, valor FROM respostas WHERE avaliacao_id=?", (st.session_state.avaliacao_id,))
                respostas_db = cursor.fetchall()
                st.session_state.respostas = {chave: valor for chave, valor in respostas_db}
                
                st.session_state.avaliacao_iniciada = True
                st.session_state.municipio = municipio_selecionado
                st.session_state.segmento = segmento_selecionado
                st.session_state.last_save_time = datetime.now()
                st.query_params["avaliacao_id"] = st.session_state.avaliacao_id
                st.rerun()

        if st.session_state.get('avaliacao_iniciada'):
            st.header(f"Avaliação: {st.session_state.get('municipio')} - {st.session_state.get('segmento')}")
            
            if 'last_save_time' not in st.session_state: st.session_state.last_save_time = datetime.now()
            if datetime.now() - st.session_state.last_save_time > timedelta(minutes=5):
                if salvar_progresso_db(conn):
                    st.session_state.last_save_time = datetime.now()
                    st.toast(f"Progresso salvo automaticamente às {datetime.now().strftime('%H:%M:%S')}")

            matriz_segmento = matriz_completa.get(st.session_state.get('segmento'), {})
            lista_secoes = list(matriz_segmento.keys())
            abas_formatadas = [s.replace('_', ' ') for s in lista_secoes]
            abas = st.tabs(abas_formatadas)

            for i, secao_atual in enumerate(lista_secoes):
                with abas[i]:
                    perguntas_da_secao = matriz_segmento[secao_atual]
                    st.subheader(f"Critérios de {abas_formatadas[i]}")
                    st.markdown("---")
                    
                    for item in perguntas_da_secao:
                        criterio = item.get("criterio", "N/A")
                        subcriterios = item.get("subcriterios", [])
                        st.markdown(f"##### {item.get('topico', 'N/A')} - {criterio}")
                        
                        st.subheader("Links de Evidência")
                        chave_links = f"{secao_atual}_{criterio}_links"
                        if chave_links not in st.session_state.respostas: st.session_state.respostas[chave_links] = "[]"
                        try: links_atuais = json.loads(st.session_state.respostas[chave_links])
                        except json.JSONDecodeError: links_atuais = []

                        for j, link in enumerate(links_atuais):
                            link_cols = st.columns([10, 1])
                            link_cols[0].info(link)
                            if link_cols[1].button("✖️", key=f"rem_{chave_links}_{j}"):
                                links_atuais.pop(j); st.session_state.respostas[chave_links] = json.dumps(links_atuais); st.rerun()
                        
                        link_cols = st.columns([10, 1])
                        novo_link = link_cols[0].text_input("Adicionar link", key=f"add_{chave_links}", label_visibility="collapsed")
                        if link_cols[1].button("➕", key=f"btn_{chave_links}"):
                            if novo_link and novo_link not in links_atuais:
                                links_atuais.append(novo_link); st.session_state.respostas[chave_links] = json.dumps(links_atuais); st.rerun()

                        st.markdown("---"); st.subheader("Critérios de Avaliação")
                        for sub in subcriterios:
                            chave_resposta = f"{secao_atual}_{criterio}_{sub}"
                            cols = st.columns([1, 2])
                            with cols[0]:
                                resposta_atual = st.session_state.respostas.get(chave_resposta, "Atende")
                                resposta = st.radio(sub, ("Atende", "Não Atende"), index=1 if resposta_atual == "Não Atende" else 0, key=chave_resposta, horizontal=True)
                                st.session_state.respostas[chave_resposta] = resposta
                            if resposta == "Não Atende":
                                with cols[1]:
                                    chave_obs = f"{chave_resposta}_obs"
                                    obs = st.text_area("Obs:", value=st.session_state.respostas.get(chave_obs, ""), key=chave_obs, label_visibility="collapsed")
                                    st.session_state.respostas[chave_obs] = obs
                        st.markdown("---")
            
            st.sidebar.header("Ações")
            if st.sidebar.button("💾 Salvar Progresso", use_container_width=True):
                if salvar_progresso_db(conn):
                    st.sidebar.success("Progresso salvo no banco de dados!")
            
            tipo_relatorio = st.sidebar.radio("Tipo de Relatório", ("Apenas Não Conformidades", "Relatório Completo"))
            
            if st.sidebar.button("📊 Gerar Relatório PDF", use_container_width=True):
                st.sidebar.info("Salvando progresso final...")
                if salvar_progresso_db(conn):
                    st.sidebar.success("Progresso salvo!")
                    with st.spinner("Gerando relatório PDF..."):
                        user_config = config['credentials']['usernames'].get(st.session_state.get('username'), {})
                        path_pdf = gerar_relatorio_novo_modelo(st.session_state.respostas, st.session_state.municipio, st.session_state.segmento, matriz_segmento, tipo_relatorio, st.session_state.get("name"), user_config)
                        st.session_state.path_pdf = path_pdf
                    if st.session_state.get('path_pdf'):
                        st.sidebar.success("Relatório PDF pronto!")
                else:
                    st.sidebar.error("Falha ao salvar. Relatório não gerado.")

            if st.session_state.get('path_pdf'):
                with open(st.session_state.path_pdf, "rb") as pdf_file:
                    st.sidebar.download_button(label="⬇️ Baixar Relatório (.pdf)", data=pdf_file, file_name=os.path.basename(st.session_state.path_pdf), mime="application/pdf")
            if st.session_state.get('fallback_docx_path'):
                with open(st.session_state.fallback_docx_path, "rb") as docx_file:
                    st.sidebar.download_button(label="⬇️ Baixar Arquivo Word (.docx)", data=docx_file, file_name=os.path.basename(st.session_state.fallback_docx_path), mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document")

    elif st.session_state.get("authentication_status") is False: st.error('Usuário ou senha incorretos.')
    elif st.session_state.get("authentication_status") is None: st.warning('Por favor, insira seu usuário e senha.')
